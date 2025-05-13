import os
import json
from pathlib import Path
import pdfplumber
from docx import Document
import chardet
from striprtf.striprtf import rtf_to_text

from preprocessing.cleaning import clean_paragraph_text

# --- Output Directories ---
text_output_dir = Path("data/processed/cases_txt")
semantic_text_output_dir = Path("data/processed/cases_vector_txt")
json_output_dir = Path("backend/data/annotated")

text_output_dir.mkdir(parents=True, exist_ok=True)
semantic_text_output_dir.mkdir(parents=True, exist_ok=True)
json_output_dir.mkdir(parents=True, exist_ok=True)

# --- Text Extraction Helpers ---
def extract_raw_text(file_path: Path) -> str:
    ext = file_path.suffix.lower()
    if ext == ".pdf":
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text.replace("\n", " ")
        return text

    elif ext == ".docx":
        doc = Document(file_path)
        return " ".join([p.text for p in doc.paragraphs])

    elif ext == ".rtf":
        with open(file_path, "rb") as f:
            raw_data = f.read()
        detected = chardet.detect(raw_data)
        encoding = detected["encoding"] or "utf-8"
        try:
            decoded_text = raw_data.decode(encoding)
            return rtf_to_text(decoded_text).replace("\n", " ")
        except UnicodeDecodeError:
            print(f"❌ Could not decode {file_path.name} with encoding: {encoding}")
            return ""
    else:
        return ""

def extract_semantic_text(file_path: Path) -> str:
    ext = file_path.suffix.lower()
    if ext == ".pdf":
        paragraphs = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    paragraphs.extend([p.strip() for p in page_text.split('\n') if p.strip()])
        return "\n\n".join(paragraphs)

    elif ext == ".docx":
        doc = Document(file_path)
        return "\n\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])

    elif ext == ".rtf":
        with open(file_path, "rb") as f:
            raw_data = f.read()
        detected = chardet.detect(raw_data)
        encoding = detected["encoding"] or "utf-8"
        try:
            decoded_text = raw_data.decode(encoding)
            text = rtf_to_text(decoded_text)
            paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
            return "\n\n".join(paragraphs)
        except UnicodeDecodeError:
            print(f"❌ Unicode decode error for {file_path.name}")
            return ""
    else:
        return ""

# --- Main callable function ---
def extract_and_clean_texts(file_path: str) -> dict:
    file_path = Path(file_path)
    case_id = file_path.stem

    raw_text = extract_raw_text(file_path)
    semantic_text = extract_semantic_text(file_path)

    if not raw_text.strip():
        raise ValueError(f"No RAW text extracted from: {file_path.name}")
    if not semantic_text.strip():
        raise ValueError(f"No SEMANTIC text extracted from: {file_path.name}")

    # Clean raw text
    cleaned_text = clean_paragraph_text(raw_text)

    # Save cleaned text
    cleaned_txt_path = text_output_dir / f"{case_id}.txt"
    with open(cleaned_txt_path, "w", encoding="utf-8") as f:
        f.write(cleaned_text)

    # Save semantic text
    semantic_txt_path = semantic_text_output_dir / f"{case_id}.txt"
    with open(semantic_txt_path, "w", encoding="utf-8") as f:
        f.write(semantic_text)

    # Update JSON
    json_path = json_output_dir / f"{case_id}.json"
    if json_path.exists():
        with open(json_path, "r", encoding="utf-8") as jf:
            data = json.load(jf)
    else:
        data = {"file_name": file_path.name}

    data["text_path"] = str(cleaned_txt_path)
    data["semantic_text_path"] = str(semantic_txt_path)

    with open(json_path, "w", encoding="utf-8") as jf:
        json.dump(data, jf, indent=2)

    return {
        "json_file": str(json_path),
        "cases_txt_path": str(cleaned_txt_path),
        "cases_vector_txt_path": str(semantic_txt_path),
    }
