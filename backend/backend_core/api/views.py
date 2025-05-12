from django.shortcuts import render, get_object_or_404
from rest_framework.decorators import api_view, parser_classes
from rest_framework.parsers import MultiPartParser, FormParser, JSONParser
from rest_framework.response import Response
from rest_framework import status
import os
import uuid
import json
import logging
from .models import Document, Annotation
import threading
from openai import OpenAI
from django.conf import settings

# Configure logging
logger = logging.getLogger(__name__)

# Initialize OpenAI client
client = OpenAI(api_key=settings.OPENAI_API_KEY)

@api_view(['GET'])
def test_api(request):
    return Response({"message": "Django API is working!"})

@api_view(['POST'])
@parser_classes([MultiPartParser, FormParser])
def upload_document(request):
    """
    Upload a document for analysis.
    """
    if 'file' not in request.FILES:
        return Response({'error': 'No file provided'}, status=status.HTTP_400_BAD_REQUEST)
    
    file = request.FILES['file']
    
    # Check file size (limit to 20MB)
    if file.size > 20 * 1024 * 1024:  # 20MB in bytes
        return Response({'error': 'File too large (max 20MB)'}, status=status.HTTP_400_BAD_REQUEST)
    
    # Check file type (accept only certain types)
    allowed_types = ['application/pdf', 'application/msword', 
                     'application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
                     'text/plain']
    
    if file.content_type not in allowed_types:
        return Response({'error': 'Invalid file type. Please upload PDF, DOC, DOCX, or TXT files.'}, 
                       status=status.HTTP_400_BAD_REQUEST)
    
    # Create document object
    document = Document(
        name=file.name,
        file=file,
        file_type=file.content_type,
        file_size=file.size,
        status='processing'
    )
    document.save()
    
    # Start asynchronous processing
    threading.Thread(target=process_document, args=(document.id,)).start()
    
    return Response({
        'documentId': str(document.id),  # Convert UUID to string
        'message': 'Document uploaded successfully and processing started'
    }, status=status.HTTP_201_CREATED)

@api_view(['GET'])
def get_document_status(request, document_id):
    """
    Get the status of a document's processing.
    """
    try:
        document = Document.objects.get(id=document_id)
        return Response({
            'documentId': str(document.id),  # Convert UUID to string
            'status': document.status
        })
    except Document.DoesNotExist:
        return Response({'error': 'Document not found'}, status=status.HTTP_404_NOT_FOUND)

@api_view(['GET'])
def get_document(request, document_id):
    """
    Get a document with its annotations.
    """
    try:
        document = Document.objects.get(id=document_id)
        annotations = document.annotations.all()
        
        # Prepare response data
        document_data = {
            'id': str(document.id),  # Convert UUID to string
            'name': document.name,
            'status': document.status,
            'content': document.content,
            'annotations': [
                {
                    'id': str(anno.id),  # Convert UUID to string
                    'text': anno.text,
                    'category': anno.category,
                    'startIndex': anno.start_index,
                    'endIndex': anno.end_index,
                    'description': anno.description
                }
                for anno in annotations
            ]
        }
        
        return Response(document_data)
    except Document.DoesNotExist:
        return Response({'error': 'Document not found'}, status=status.HTTP_404_NOT_FOUND)

@api_view(['POST'])
def chat_with_document(request):
    """
    Chat with a document using the OpenAI API.
    """
    data = request.data
    message = data.get('message')
    selected_text = data.get('selectedText')
    document_id = data.get('documentId')
    history = data.get('history', [])
    
    if not message:
        return Response({'error': 'No message provided'}, status=status.HTTP_400_BAD_REQUEST)
    
    try:
        # If document_id is provided, get the document
        document_context = ""
        selected_context = ""
        if document_id:
            try:
                document = Document.objects.get(id=document_id)
                document_context = document.content
                
                # If selected_text is provided, add context about it
                if selected_text:
                    # Find matching annotation if possible
                    annotations = document.annotations.filter(text=selected_text)
                    if annotations.exists():
                        annotation = annotations.first()
                        selected_context = f"Selected text: '{selected_text}' (Category: {annotation.category})"
                        if annotation.description:
                            selected_context += f", Description: {annotation.description}"
                    else:
                        selected_context = f"Selected text: '{selected_text}'"
            except Document.DoesNotExist:
                return Response({'error': 'Document not found'}, status=status.HTTP_404_NOT_FOUND)
        
        # Construct messages for OpenAI API
        messages = [
            {"role": "system", "content": "You are a legal assistant specializing in document analysis. Provide clear, concise, and accurate responses to questions about legal documents."}
        ]
        
        # Add document context if available
        if document_context:
            # Truncate to prevent hitting token limits
            truncated_context = document_context[:4000] + ("..." if len(document_context) > 4000 else "")
            messages.append({"role": "system", "content": f"The following is the content of a legal document: {truncated_context}"})
        
        # Add selected text context if available
        if selected_context:
            messages.append({"role": "system", "content": selected_context})
        
        # Add conversation history
        for msg in history:
            role = "assistant" if msg.get('role') == 'assistant' else "user"
            messages.append({"role": role, "content": msg.get('content', '')})
        
        # Add the current user message
        messages.append({"role": "user", "content": message})
        
        # Call OpenAI API
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",  # Adjust based on your needs
            messages=messages,
            max_tokens=500,
            temperature=0.7
        )
        
        # Extract response text
        ai_response = response.choices[0].message.content
        
        # Prepare sources (for demo, use mock sources)
        sources = []
        
        return Response({
            'response': ai_response,
            'sources': sources
        })
    
    except Exception as e:
        logger.error(f"Error in chat_with_document: {str(e)}")
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

def process_document(document_id):
    """
    Process a document to extract its content and add annotations.
    This runs in a separate thread.
    """
    try:
        document = Document.objects.get(id=document_id)
        document.status = 'processing'
        document.save()
        
        # Extract content from file (simplified)
        file_path = document.file.path
        content = ""
        
        # Simple content extraction based on file type
        if document.file_type == 'text/plain':
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
        elif document.file_type.startswith('application/pdf'):
            # Use a PDF parser library in a real application
            try:
                import PyPDF2
                with open(file_path, 'rb') as pdf_file:
                    reader = PyPDF2.PdfReader(pdf_file)
                    content = ""
                    for page_num in range(len(reader.pages)):
                        content += reader.pages[page_num].extract_text() + "\n"
            except Exception as e:
                logger.error(f"Error extracting PDF content: {str(e)}")
                content = "Error extracting PDF content. Please try another document."
        elif document.file_type.startswith('application/'):
            # Use appropriate document parser for other formats
            try:
                import docx
                doc = docx.Document(file_path)
                content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            except Exception as e:
                logger.error(f"Error extracting DOCX content: {str(e)}")
                content = "Error extracting document content. Please try another document."
        else:
            content = "Document content could not be extracted. Unsupported format."
        
        document.content = content
        document.save()
        
        # Analyze document with OpenAI to extract annotations
        analyze_document_with_openai(document)
        
        document.status = 'analyzed'
        document.save()
        
    except Exception as e:
        logger.error(f"Error processing document {document_id}: {str(e)}")
        try:
            document = Document.objects.get(id=document_id)
            document.status = 'error'
            document.save()
        except:
            pass

def analyze_document_with_openai(document):
    """
    Use OpenAI to analyze the document and extract annotations.
    """
    try:
        # Truncate content if it's too long
        content = document.content[:5000]  # Adjust based on token limits
        
        # Create a prompt for OpenAI
        prompt = f"""
        Analyze the following legal document and identify key elements in these categories:
        1. Legal terms
        2. Dates
        3. Parties involved
        4. Obligations
        5. Conditions
        
        For each element, provide:
        - The exact text snippet
        - The category it belongs to
        - A brief description or explanation
        - The starting and ending character positions in the text (if possible)
        
        Document content:
        {content}
        """
        
        # Call OpenAI API
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",  # Adjust based on your needs
            messages=[
                {"role": "system", "content": "You are a legal document analysis assistant."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1000,
            temperature=0.3
        )
        
        # Extract the response
        analysis_text = response.choices[0].message.content
        
        # Parse the analysis text and create annotations
        # This is a simplified version; in practice you'd need more robust parsing
        lines = analysis_text.split('\n')
        current_category = None
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Check if line defines a category
            if line.lower().startswith('legal term'):
                current_category = 'legal_term'
            elif line.lower().startswith('date'):
                current_category = 'date'
            elif line.lower().startswith('part'):
                current_category = 'party'
            elif line.lower().startswith('obligation'):
                current_category = 'obligation'
            elif line.lower().startswith('condition'):
                current_category = 'condition'
            elif ':' in line and current_category:
                # This might be an annotation
                parts = line.split(':', 1)
                if len(parts) == 2:
                    text = parts[1].strip()
                    description = None
                    
                    # Simple position estimation (in a real app, use more precise methods)
                    if text in document.content:
                        start_index = document.content.find(text)
                        end_index = start_index + len(text)
                        
                        # Create annotation
                        Annotation.objects.create(
                            document=document,
                            text=text,
                            category=current_category,
                            start_index=start_index,
                            end_index=end_index,
                            description=description
                        )
        
        # Create some sample annotations for testing if none were created
        if document.annotations.count() == 0:
            sample_texts = [
                ("Contract", "legal_term", "A legally binding agreement"),
                ("January 1, 2023", "date", "Effective date of the agreement"),
                ("John Smith", "party", "First party to the agreement"),
                ("must pay within 30 days", "obligation", "Payment timeline"),
                ("unless terminated earlier", "condition", "Early termination condition")
            ]
            
            for sample_text, category, description in sample_texts:
                if sample_text in document.content:
                    start_index = document.content.find(sample_text)
                    end_index = start_index + len(sample_text)
                else:
                    # Just place it at the beginning if not found
                    start_index = 0
                    end_index = len(sample_text)
                
                Annotation.objects.create(
                    document=document,
                    text=sample_text,
                    category=category,
                    start_index=start_index,
                    end_index=end_index,
                    description=description
                )
                
    except Exception as e:
        logger.error(f"Error analyzing document with OpenAI: {str(e)}")
        # Create one basic annotation to show something on the frontend
        Annotation.objects.create(
            document=document,
            text="Error occurred during analysis",
            category="other",
            start_index=0,
            end_index=len("Error occurred during analysis"),
            description="There was an error analyzing this document"
        )

@api_view(['GET'])
def debug_info(request):
    """
    Debug endpoint to verify API is working correctly
    """
    debug_data = {
        'api_version': '1.0.0',
        'openai_configured': bool(settings.OPENAI_API_KEY),
        'debug_mode': settings.DEBUG,
        'allowed_hosts': settings.ALLOWED_HOSTS,
        'media_root': settings.MEDIA_ROOT,
        'cors_allowed': getattr(settings, 'CORS_ALLOW_ALL_ORIGINS', False)
    }
    return Response(debug_data)
