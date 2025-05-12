import os
import re
import json
from pathlib import Path
import pdfplumber
from docx import Document
import chardet
from striprtf.striprtf import rtf_to_text

# --- Helper functions to extract text from supported file types ---
def extract_raw_text(file_path: Path) -> str:
    ext = file_path.suffix.lower()

    if ext == ".pdf":
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text.replace("\n", " ")
        return text

    elif ext == ".docx":
        doc = Document(file_path)
        return " ".join([p.text for p in doc.paragraphs])

    elif ext == ".rtf":
    # Read the raw bytes
        with open(file_path, "rb") as f:
            raw_data = f.read()

        # Detect encoding
        detected = chardet.detect(raw_data)
        encoding = detected["encoding"] or "utf-8"  # fallback to utf-8 just in case

        try:
            decoded_text = raw_data.decode(encoding)
            return rtf_to_text(decoded_text).replace("\n", " ")
        except UnicodeDecodeError:
            print(f"❌ Still failed to decode {file_path.name} with detected encoding: {encoding}")
            return ""

    else:
        return ""

def extract_semantic_text(file_path: Path) -> str:
    """Extracts text keeping better paragraph structure for semantic search."""
    ext = file_path.suffix.lower()

    if ext == ".pdf":
        paragraphs = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    page_paragraphs = [p.strip() for p in page_text.split('\n') if p.strip()]
                    paragraphs.extend(page_paragraphs)
        return "\n\n".join(paragraphs)

    elif ext == ".docx":
        doc = Document(file_path)
        return "\n\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])

    elif ext == ".rtf":
        with open(file_path, "rb") as f:
            raw_data = f.read()
        detected = chardet.detect(raw_data)
        encoding = detected["encoding"] or "utf-8"
        try:
            decoded_text = raw_data.decode(encoding)
            text = rtf_to_text(decoded_text)
            paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
            return "\n\n".join(paragraphs)
        except UnicodeDecodeError:
            print(f"❌ Unicode decode error for {file_path.name}")
            return ""

    else:
        return ""

# --- Paths ---
documents_dir = Path("data/raw/Documents")
json_output_dir = Path("backend/data/annotated")
text_output_dir = Path("data/processed/cases_txt")
semantic_text_output_dir = Path("data/processed/cases_vector_txt")

text_output_dir.mkdir(parents=True, exist_ok=True)
semantic_text_output_dir.mkdir(parents=True, exist_ok=True)

# --- Main Loop ---
for file in documents_dir.iterdir():
    if file.suffix.lower() in [".pdf", ".docx", ".rtf"]:
        case_id = file.stem

        # Extract texts
        raw_text = extract_raw_text(file)
        semantic_text = extract_semantic_text(file)

        if not raw_text.strip():
            print(f"⚠️ No RAW text extracted from: {file.name}")
            continue
        if not semantic_text.strip():
            print(f"⚠️ No SEMANTIC text extracted from: {file.name}")
            continue

        # Save raw text
        raw_txt_path = text_output_dir / f"{case_id}.txt"
        with open(raw_txt_path, "w", encoding="utf-8") as out_f:
            out_f.write(raw_text)

        # Save semantic text
        semantic_txt_path = semantic_text_output_dir / f"{case_id}.txt"
        with open(semantic_txt_path, "w", encoding="utf-8") as out_f:
            out_f.write(semantic_text)

        # Update JSON
        json_file_path = json_output_dir / f"{case_id}.json"
        if json_file_path.exists():
            with open(json_file_path, "r", encoding="utf-8") as jf:
                data = json.load(jf)

            # Make sure we keep the original `text_path`
            if "text_path" not in data:
                data["text_path"] = str(raw_txt_path)

            # Add new semantic text path
            data["semantic_text_path"] = str(semantic_txt_path)

            with open(json_file_path, "w", encoding="utf-8") as jf:
                json.dump(data, jf, indent=2)

print("✅ All supported files processed. Raw and semantic text paths updated in JSONs.")